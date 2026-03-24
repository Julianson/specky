"""
DOCX Template — Microsoft-Branded Word Document Generator
Uses python-docx to create professional .docx files with Microsoft branding.

Usage:
    python3 docx_template.py

Requirements:
    pip install python-docx

Customization:
    Replace the {{variables}} and sample content with real data.
    The agent should modify this template for each document request.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date
import os

# === CONFIG ===
TITLE = "{{title}}"
SUBTITLE = "{{subtitle}}"
AUTHOR = "{{author_name}}"
CLIENT = "{{client_name}}"
VERSION = "1.0.0"
DATE = date.today().isoformat()

# === MICROSOFT BRAND COLORS ===
BLUE = RGBColor(0x00, 0x78, 0xD4)       # #0078D4 — Primary
RED = RGBColor(0xF2, 0x50, 0x22)         # #F25022
GREEN = RGBColor(0x7F, 0xBA, 0x00)       # #7FBA00
LIGHT_BLUE = RGBColor(0x00, 0xA4, 0xEF)  # #00A4EF
YELLOW = RGBColor(0xFF, 0xB9, 0x00)      # #FFB900
TEXT_PRIMARY = RGBColor(0x32, 0x31, 0x30) # #323130
TEXT_SECONDARY = RGBColor(0x60, 0x5E, 0x5C) # #605E5C
ALT_ROW = RGBColor(0xF3, 0xF2, 0xF1)    # #F3F2F1

# 4-color rotation for section headings
SECTION_COLORS = [RED, GREEN, LIGHT_BLUE, YELLOW]


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def add_4color_bar(doc):
    """Add the Microsoft 4-color bar as a thin table."""
    colors = ['F25022', '7FBA00', '00A4EF', 'FFB900']
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, color in enumerate(colors):
        cell = table.cell(0, i)
        cell.width = Inches(1.625)
        set_cell_shading(cell, color)
        cell.text = ''
        for p in cell.paragraphs:
            p.space_before = Pt(0)
            p.space_after = Pt(0)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
    # Make rows very thin
    for row in table.rows:
        row.height = Pt(4)


def create_cover_page(doc):
    """Create the branded cover page."""
    # 4-color bar at top
    add_4color_bar(doc)

    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(TITLE)
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE
    run.font.bold = True
    run.font.name = 'Segoe UI'

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(SUBTITLE)
    run.font.size = Pt(16)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = 'Segoe UI'

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ('Author', AUTHOR),
        ('Date', DATE),
        ('Version', VERSION),
        ('Client', CLIENT),
        ('Status', 'Draft'),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.cell(i, 0).text = label
        meta_table.cell(i, 1).text = value
        for cell in meta_table.row_cells(i):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10)

    # Confidentiality
    doc.add_paragraph()
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf_para.add_run('Microsoft Confidential')
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = 'Segoe UI'
    run.font.italic = True

    # 4-color bar at bottom
    add_4color_bar(doc)

    doc.add_page_break()


def add_heading(doc, text, level=1, color=None):
    """Add a branded heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Segoe UI Semibold'
        if color:
            run.font.color.rgb = color
        elif level == 1:
            run.font.color.rgb = BLUE


def add_branded_table(doc, headers, rows):
    """Create a table with Microsoft-branded styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '0078D4')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows with alternating colors
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = str(value)
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'F3F2F1')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(9)
                    run.font.color.rgb = TEXT_PRIMARY

    return table


def add_callout(doc, text, callout_type='info'):
    """Add a callout/highlight box."""
    colors_map = {
        'info': '0078D4',
        'warning': 'FFB900',
        'success': '107C10',
        'critical': 'E81123',
    }
    color = colors_map.get(callout_type, '0078D4')
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_shading(cell, color + '20')  # Light tint
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9)


def create_document():
    """Generate the complete branded document."""
    doc = Document()

    # === COVER PAGE ===
    create_cover_page(doc)

    # === DOCUMENT HISTORY ===
    add_heading(doc, 'Document History', level=2)
    add_branded_table(doc, ['Version', 'Date', 'Author', 'Changes'], [
        [VERSION, DATE, AUTHOR, 'Initial version'],
    ])
    doc.add_page_break()

    # === TABLE OF CONTENTS ===
    add_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph('(Auto-generated — update field after editing)')
    doc.add_page_break()

    # === EXECUTIVE SUMMARY ===
    add_heading(doc, 'Executive Summary', level=1)
    doc.add_paragraph(
        'This document provides an overview of... '
        '[Replace with 150-250 word executive summary]'
    )

    # === MAIN CONTENT SECTIONS ===
    sections = [
        ('Section 1: Overview', 'Content for section 1...'),
        ('Section 2: Analysis', 'Content for section 2...'),
        ('Section 3: Recommendations', 'Content for section 3...'),
    ]
    for i, (title, content) in enumerate(sections):
        color = SECTION_COLORS[i % len(SECTION_COLORS)]
        add_heading(doc, title, level=1, color=color)
        doc.add_paragraph(content)

    # Example table in a section
    add_heading(doc, 'Key Metrics', level=2)
    add_branded_table(doc, ['Metric', 'Current', 'Target', 'Status'], [
        ['Metric 1', '75%', '90%', 'In Progress'],
        ['Metric 2', '60%', '85%', 'At Risk'],
        ['Metric 3', '95%', '95%', 'On Track'],
    ])

    # === NEXT STEPS ===
    doc.add_page_break()
    add_heading(doc, 'Next Steps', level=1)
    for step in ['Action item 1 — Owner — Deadline', 'Action item 2 — Owner — Deadline']:
        doc.add_paragraph(step, style='List Bullet')

    # === REFERENCES ===
    add_heading(doc, 'References', level=1)
    doc.add_paragraph('[1] Source URL — Description')

    # === SAVE ===
    os.makedirs('output/docx', exist_ok=True)
    filename = f"{TITLE.replace(' ', '_')}_v{VERSION}_{DATE}.docx"
    filepath = f"output/docx/{filename}"
    doc.save(filepath)
    print(f"Created: {filepath}")


if __name__ == '__main__':
    create_document()
